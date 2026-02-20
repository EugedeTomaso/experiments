import io
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .base import BaseExporter
from ..export_utils import build_combined_markdown


class DocxExporter(BaseExporter):
    @property
    def content_type(self):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self):
        return "docx"

    def export(self) -> bytes:
        if self.nodes_with_depth:
            raw_md = build_combined_markdown(self.nodes_with_depth)
        else:
            raw_md = self.content_md

        doc = Document()

        # Title
        title_para = doc.add_heading(self.title, level=0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Parse markdown lines into docx elements
        self._md_to_docx(doc, raw_md)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _md_to_docx(self, doc, md_text):
        """Simple markdown to docx converter handling headings, paragraphs, bold, italic."""
        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                doc.add_heading(heading_match.group(2), level=level)
                i += 1
                continue

            # Blank line
            if not line.strip():
                i += 1
                continue

            # Regular paragraph — collect consecutive non-blank, non-heading lines
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r"^#{1,6}\s+", lines[i]):
                para_lines.append(lines[i])
                i += 1

            text = " ".join(para_lines)
            para = doc.add_paragraph()
            self._add_formatted_runs(para, text)

    def _add_formatted_runs(self, para, text):
        """Parse inline markdown (bold, italic) into formatted runs."""
        # Split on **bold** and *italic* patterns
        parts = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", text)
        for part in parts:
            if part.startswith("***") and part.endswith("***"):
                run = para.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)
